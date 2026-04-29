"""Generate a formatted Word document from ROLE_USE_CASES.md"""
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Default font --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

# Heading styles
for level, (size, color) in {1: (22, '1a3764'), 2: (16, '1a5276'), 3: (13, '2c3e50'), 4: (11, '34495e')}.items():
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor(*(int(color[i:i+2], 16) for i in (0, 2, 4)))
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    hs.paragraph_format.space_after = Pt(6)

BRAND_BLUE = RGBColor(0x1a, 0x37, 0x64)
BRAND_LIGHT = RGBColor(0xd6, 0xe4, 0xf0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_BG = RGBColor(0xf2, 0xf2, 0xf2)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(headers, rows, col_widths=None, header_color='1a3764'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            # Handle checkmarks and crosses
            if val.strip() in ('✅', '❌'):
                run = p.add_run(val.strip())
                run.font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Handle bold markers
                parts = re.split(r'\*\*(.+?)\*\*', val)
                for j, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
                    if j % 2 == 1:
                        run.bold = True
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'f7f9fc')

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

# ============== COVER PAGE ==============
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('Q-CRM')
run.font.size = Pt(36)
run.font.color.rgb = BRAND_BLUE
run.bold = True
run.font.name = 'Calibri'

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run('Role Use Cases & Permission Guide')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'Calibri'

doc.add_paragraph()

line_p = doc.add_paragraph()
line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = line_p.add_run('━' * 40)
run.font.color.rgb = BRAND_BLUE
run.font.size = Pt(12)

doc.add_paragraph()

for label, value in [
    ('Application: ', 'https://qcrm.qbadvisory.com'),
    ('Date: ', 'April 24, 2026'),
    ('Version: ', '1.0'),
    ('Classification: ', 'Internal'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = 'Calibri'
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.bold = True
    run.font.name = 'Calibri'

doc.add_page_break()

# ============== TABLE OF CONTENTS ==============
doc.add_heading('Table of Contents', level=1)

toc_items = [
    '1. Opportunity Lifecycle Workflow',
    '2. Role Overview',
    '3. Admin',
    '4. Manager',
    '5. Sales',
    '6. Presales',
    '7. Management',
    '8. Read-Only',
    '9. Permission Reference Matrix',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND_BLUE
    run.font.name = 'Calibri'

doc.add_page_break()

# ============== 1. OPPORTUNITY LIFECYCLE WORKFLOW ==============
doc.add_heading('1. Opportunity Lifecycle Workflow', level=1)

p = doc.add_paragraph()
run = p.add_run('The Q-CRM opportunity lifecycle flows through six pipeline stages. Each stage has a probability weight, required fields, SLA timers, and role-specific actions. The diagram below shows the complete flow from lead intake through closure.')
run.font.size = Pt(10)
run.font.name = 'Calibri'

# -- 1.1 Pipeline Stage Flow (visual table-based diagram) --
doc.add_heading('1.1 Pipeline Stage Flow', level=2)

# Create a visual flow using a single-row table with arrows
flow_table = doc.add_table(rows=3, cols=11)
flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Stage names and probabilities
stages_flow = [
    ('Discovery', '10%'),
    ('→', ''),
    ('Qualification', '25%'),
    ('→', ''),
    ('Proposal', '50%'),
    ('→', ''),
    ('Negotiation', '75%'),
    ('→', ''),
    ('Closed Won', '100%'),
]

# Row 0: Stage names
for i, (name, _) in enumerate(stages_flow):
    cell = flow_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.name = 'Calibri'
    if name == '→':
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    else:
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, '1a3764')

# Row 1: Probabilities
for i, (name, prob) in enumerate(stages_flow):
    cell = flow_table.rows[1].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(prob)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Calibri'

# Row 2: empty spacer
for i in range(11):
    flow_table.rows[2].cells[i].text = ''

# Closed Lost note
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('↳  Closed Lost (0%)')
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
run.font.name = 'Calibri'
run = p.add_run('  — can happen from any stage')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'

# Re-estimate loop note
p = doc.add_paragraph()
run = p.add_run('↺  Re-estimate Loop')
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)
run.font.name = 'Calibri'
run = p.add_run('  — Sales can send back from Proposal or Negotiation to Qualification (resets GOM approval)')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'

doc.add_paragraph()

# -- 1.2 Back-and-Forth Workflows --
doc.add_heading('1.2 Back-and-Forth Workflows', level=2)

p = doc.add_paragraph()
run = p.add_run('Q-CRM is not a simple linear pipeline. Multiple feedback loops drive opportunities backward through stages. Below are all bidirectional workflows in the system.')
run.font.size = Pt(10)
run.font.name = 'Calibri'

# --- 1.2.1 Re-Estimate Loop ---
doc.add_heading('1.2.1 Re-Estimate Loop (Sales → Presales)', level=3)

p = doc.add_paragraph()
run = p.add_run('Sales can send an opportunity back from Proposal or Negotiation to Qualification for re-estimation. This is the most common feedback loop.')
run.font.size = Pt(10)
run.font.name = 'Calibri'

# Visual flow table for re-estimate
reest_flow = doc.add_table(rows=1, cols=5)
reest_flow.alignment = WD_TABLE_ALIGNMENT.CENTER
reest_flow.style = 'Table Grid'
flow_data = [
    ('PROPOSAL\nor\nNEGOTIATION', '1a3764'),
    ('→ Send for\nRe-estimate →', None),
    ('QUALIFICATION\n(GOM resets)', 'c0392b'),
    ('→ Re-submit\nestimation →', None),
    ('PROPOSAL\n(return)', '27ae60'),
]
for i, (text, bg) in enumerate(flow_data):
    cell = reest_flow.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    run.bold = True
    if bg:
        run.font.color.rgb = WHITE
        set_cell_shading(cell, bg)
    else:
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.bold = False

doc.add_paragraph()

reest_rows = [
    ['Trigger',         'Sales clicks "Send for Re-estimate" from Proposal or Negotiation stage'],
    ['Stage Change',    'Opportunity moves back to Qualification'],
    ['GOM Reset',       'gomApproved is reset to FALSE — must pass approval gate again'],
    ['Counter',         'reEstimateCount is incremented (+1 each loop)'],
    ['Status Update',   'detailedStatus = "Sent for Re-estimate"'],
    ['Audit Trail',     'SEND_BACK_REESTIMATE audit entry created with optional comment'],
    ['Email Trigger',   '"sent_back_to_reestimate" email sent to Presales and Manager'],
    ['Return Path',     'Presales re-works → GOM recalculated → must pass approval gate → resubmits → detailedStatus = "Re-estimation Submitted" → back to Proposal'],
]

reest_table = doc.add_table(rows=len(reest_rows), cols=2)
reest_table.alignment = WD_TABLE_ALIGNMENT.LEFT
reest_table.style = 'Table Grid'
for r_idx, (label, desc) in enumerate(reest_rows):
    cell0 = reest_table.rows[r_idx].cells[0]
    cell0.text = ''
    p = cell0.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = WHITE
    set_cell_shading(cell0, '2c3e50')
    cell0.width = Cm(3)

    cell1 = reest_table.rows[r_idx].cells[1]
    cell1.text = ''
    p = cell1.paragraphs[0]
    run = p.add_run(desc)
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    if r_idx % 2 == 1:
        set_cell_shading(cell1, 'f7f9fc')

doc.add_paragraph()

# --- 1.2.2 GOM Approval Cycle ---
doc.add_heading('1.2.2 GOM Approval Cycle (Presales ↔ Manager)', level=3)

p = doc.add_paragraph()
run = p.add_run('Within the Qualification stage, the GOM approval can loop between Presales and Manager until the margin is acceptable:')
run.font.size = Pt(10)
run.font.name = 'Calibri'

# Visual flow
gom_loop = doc.add_table(rows=2, cols=5)
gom_loop.alignment = WD_TABLE_ALIGNMENT.CENTER
gom_loop.style = 'Table Grid'
gom_flow = [
    ('Presales\ncalculates GOM', '1a3764'),
    ('→ submit →', None),
    ('Manager\nreviews', '1a5276'),
    ('→ rejects →', None),
    ('Presales\nrevises & \nresubmits', 'e67e22'),
]
for i, (text, bg) in enumerate(gom_flow):
    cell = gom_loop.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    run.bold = True
    if bg:
        run.font.color.rgb = WHITE
        set_cell_shading(cell, bg)
    else:
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.bold = False
# Row 2: approval exit
for i in range(5):
    cell = gom_loop.rows[1].cells[i]
    cell.text = ''
if True:
    cell = gom_loop.rows[1].cells[2]
    p = cell.paragraphs[0]
    run = p.add_run('→ approves →\ngomApproved = true\n→ PROPOSAL')
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xae, 0x60)

doc.add_paragraph()

gom_gate_table = doc.add_table(rows=5, cols=3)
gom_gate_table.alignment = WD_TABLE_ALIGNMENT.LEFT
gom_gate_table.style = 'Table Grid'

gom_headers = ['Condition', 'Action', 'Result']
for i, h in enumerate(gom_headers):
    cell = gom_gate_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'
    set_cell_shading(cell, '1a3764')

gom_rows = [
    ['GOM% ≥ auto-approve threshold', 'Auto-approved by system', 'gomApproved = true ✅\nCan proceed to Proposal'],
    ['GOM% < threshold', 'ApprovalRequest created', 'Routed to Reporting Manager\nAwaits manual review'],
    ['Manager approves', 'GOM approval granted', 'gomApproved = true ✅\nCan proceed to Proposal'],
    ['Manager rejects', 'GOM approval denied', 'gomApproved = false ❌\nPresales revises → resubmits (LOOP)'],
]
for r_idx, row_data in enumerate(gom_rows):
    for c_idx, val in enumerate(row_data):
        cell = gom_gate_table.rows[r_idx + 1].cells[c_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        if r_idx % 2 == 1:
            set_cell_shading(cell, 'f7f9fc')

doc.add_paragraph()

# --- 1.2.3 Discount Approval Cycle ---
doc.add_heading('1.2.3 Discount Approval Cycle (Sales ↔ Finance)', level=3)

p = doc.add_paragraph()
run = p.add_run('When Sales sets a discount above the threshold, it triggers a Finance approval loop:')
run.font.size = Pt(10)
run.font.name = 'Calibri'

disc_flow = doc.add_table(rows=1, cols=5)
disc_flow.alignment = WD_TABLE_ALIGNMENT.CENTER
disc_flow.style = 'Table Grid'
disc_data = [
    ('Sales sets\ndiscount > 15%\n& margin < 20%', '1a3764'),
    ('→ approval\nrequest →', None),
    ('Finance Mgr\nreviews', '1a5276'),
    ('→ rejects →', None),
    ('Sales adjusts\npricing &\nresubmits', 'e67e22'),
]
for i, (text, bg) in enumerate(disc_data):
    cell = disc_flow.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    run.bold = True
    if bg:
        run.font.color.rgb = WHITE
        set_cell_shading(cell, bg)
    else:
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.bold = False

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Auto-approve: ')
run.bold = True
run.font.size = Pt(9)
run.font.name = 'Calibri'
run = p.add_run('If discount ≤ 15% OR margin ≥ 20%, the discount is auto-approved without Finance review.')
run.font.size = Pt(9)
run.font.name = 'Calibri'

doc.add_paragraph()

# --- 1.2.4 SOW Approval Chain ---
doc.add_heading('1.2.4 SOW Approval Chain (Multi-Step Review)', level=3)

p = doc.add_paragraph()
run = p.add_run('SOW approval uses a configurable multi-step chain. Rejection at any step sends the SOW back to Drafting, and the entire chain restarts:')
run.font.size = Pt(10)
run.font.name = 'Calibri'

sow_chain = doc.add_table(rows=3, cols=5)
sow_chain.alignment = WD_TABLE_ALIGNMENT.CENTER
sow_chain.style = 'Table Grid'
chain_steps = ['Sales\nReview', 'Presales\nReview', 'Manager\nApproval', 'Finance\nReview', 'Legal\nReview']

for i, step in enumerate(chain_steps):
    cell = sow_chain.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Step {i+1}')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Calibri'

    cell = sow_chain.rows[1].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(step)
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.bold = True
    run.font.color.rgb = WHITE
    set_cell_shading(cell, '1a3764')

    cell = sow_chain.rows[2].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Approve ✅\nor Reject ❌')
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('On REJECT at any step: ')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
run.font.name = 'Calibri'
run = p.add_run('SOW goes back to DRAFTING → Author fixes → resubmits → approval chain restarts from Step 1. Each step has configurable reviewerRoles, escalationHours, and isRequired flag.')
run.font.size = Pt(9)
run.font.name = 'Calibri'

doc.add_paragraph()

# --- 1.2.5 SOW Client Revision Loop ---
doc.add_heading('1.2.5 SOW Document Revision Loop (Team ↔ Client)', level=3)

sow_loop_flow = doc.add_table(rows=1, cols=7)
sow_loop_flow.alignment = WD_TABLE_ALIGNMENT.CENTER
sow_loop_flow.style = 'Table Grid'
sow_flow_data = [
    ('DRAFTING', '1a3764'),
    ('→', None),
    ('IN REVIEW\n& APPROVED', '1a5276'),
    ('→', None),
    ('SHARED WITH\nCLIENT', '2c3e50'),
    ('→ revision\nrequested →', None),
    ('DRAFTING\n(loop back)', 'e67e22'),
]
for i, (text, bg) in enumerate(sow_flow_data):
    cell = sow_loop_flow.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    run.bold = True
    if bg:
        run.font.color.rgb = WHITE
        set_cell_shading(cell, bg)
    else:
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.bold = False

doc.add_paragraph()

sow_lifecycle = [
    ['1',  'Not Started',          'Document created, no content yet'],
    ['2',  'Drafting',             'User has started editing sections (← REVISION loops back here)'],
    ['3',  'AI Generated',        'AI has generated section content'],
    ['4',  'In Review',           'Submitted for internal review'],
    ['5',  'Requires Inputs',     'Missing data flagged by readiness engine (← loops back to Drafting)'],
    ['6',  'Approved Internally', 'All internal approval steps passed'],
    ['7',  'Shared with Client',  'Sent to client for review'],
    ['8',  'Revision Requested',  'Client requested changes (← loops back to Drafting)'],
    ['9',  'Finalized',           'All parties agreed on terms'],
    ['10', 'Signed / Archived',   'Executed and archived — end state'],
]

add_table(
    ['#', 'Status', 'Description'],
    sow_lifecycle,
    col_widths=[1, 3.5, 10.5]
)

# -- 1.3 Complete Handoff Summary --
doc.add_heading('1.3 Complete Back-and-Forth Summary', level=2)

p = doc.add_paragraph()
run = p.add_run('All bidirectional workflows in Q-CRM at a glance:')
run.font.size = Pt(10)
run.font.name = 'Calibri'

handoff_rows = [
    ['1', 'Re-estimate', 'Proposal/Negotiation\n→ Qualification', 'Sales clicks\n"Send for Re-estimate"', 'gomApproved=false\ndetailedStatus reset\nreEstimateCount+1', 'Presales re-works →\nresubmits → Proposal'],
    ['2', 'GOM rejection', 'Manager → Presales\n(within Qualification)', 'Manager rejects\nGOM approval', 'ApprovalRequest\nstatus=Rejected', 'Presales revises →\nresubmits GOM'],
    ['3', 'Discount rejection', 'Finance → Sales\n(within Negotiation)', 'Finance rejects\ndiscount request', 'ApprovalRequest\nstatus=Rejected', 'Sales adjusts\npricing → resubmits'],
    ['4', 'SOW review\nrejection', 'Any reviewer →\nAuthor (within SOW)', 'Reviewer rejects\nSOW step', 'SOW status reverts\nto Drafting', 'Author fixes →\nresubmits → chain\nrestarts'],
    ['5', 'SOW client\nrevision', 'Client → Author\n(within SOW)', 'Client requests\nchanges', 'SOW status =\nRevision Requested', 'Author revises →\nre-shares with client'],
    ['6', 'Requires Inputs', 'Readiness engine →\nAuthor (within SOW)', 'Missing data detected\nduring review', 'SOW status =\nRequires Inputs', 'Author fills gaps →\nresubmits for review'],
]

add_table(
    ['#', 'Loop Name', 'Direction', 'Trigger', 'What Resets', 'Return Path'],
    handoff_rows,
    col_widths=[0.6, 2, 2.5, 2.5, 2.5, 3]
)

# -- 1.4 Role Participation by Stage --
doc.add_heading('1.4 Role Participation by Stage', level=2)

role_participation = [
    ['Lead Intake',     'Creates',         '—',                '—',                'Notified', '—',              '—'],
    ['Discovery',       'Fills pipeline',  '—',                'Notified',         'Notified', '—',              'Views'],
    ['Qualification',   'Reviews',         '**Estimates & GOM**', '**Approves GOM**', '—',       '—',              'Views'],
    ['Proposal',        '**Drives proposal**', 'Views',        'Notified',         '—',        '—',              'Views'],
    ['Negotiation',     '**Negotiates**',  'Views',            'Oversight',        'Notified', 'Views',          'Views'],
    ['Closed Won',      '**Converts**',    'Notified',         'Notified',         'Notified', 'Views analytics','Views'],
    ['Closed Lost',     'Records reason',  '—',                'Notified',         'Notified', 'Views analytics','Views'],
]

add_table(
    ['Stage', 'Sales', 'Presales', 'Manager', 'Admin', 'Management', 'Read-Only'],
    role_participation,
    col_widths=[2.2, 2.2, 2.2, 2.2, 1.8, 2.2, 1.8]
)

# -- 1.5 Detailed Lifecycle Steps --
doc.add_heading('1.5 Detailed Lifecycle Steps', level=2)

lifecycle_steps = [
    ['1', 'Lead Intake', 'External system, manual entry, or chatbot', 'Auto-scoring (Title + Budget + Company + Source)\nAuto-creates Client & Contact, deduplicates within 60 days\nScore > 70 = Hot, 40–70 = Warm, < 40 = Low fit', 'Sales'],
    ['2', 'Discovery (10%)', 'Sales fills pipeline data', 'Title, client, value, region, practice, technology, pricing model, day rate, dates\nTriggers "moved_to_presales" email', 'Sales creates\nManager notified'],
    ['3', 'Qualification (25%)', 'Presales builds estimation', 'Resource plan, effort estimates, cost model via rate cards\nGOM Calculator: adjustedCost = CTC × (1 + overheads)\nGOM% = (Revenue − Cost) / Revenue × 100\n\n⚠ GOM APPROVAL GATE — must pass before Proposal', 'Presales estimates\nManager approves GOM'],
    ['4', 'Proposal (50%)', 'Sales drives proposal', 'Triggers "presales_submitted_back" email\nSOW document generation can begin\nCan be sent back to Qualification (re-estimate)', 'Sales drives\nPresales views'],
    ['5', 'Negotiation (75%)', 'Commercial negotiation', 'Contract finalization and discount approval\nDiscount > 15% AND margin < 20% → Finance approval\nCan be sent back to Qualification (re-estimate)', 'Sales negotiates\nManager oversight'],
    ['6a', 'Closed Won (100%)', 'Deal won', 'actualCloseDate set, "proposal_won" email to all\nDeal → Project conversion available\nSets detailedStatus = "SOW Approved"', 'Sales converts\nAll notified'],
    ['6b', 'Closed Lost (0%)', 'Deal lost', 'actualCloseDate set, lostRemarks captured\nMARK_LOST audit entry\n"proposal_lost" email to Admin/Manager', 'Sales records\nAdmin/Mgr notified'],
]

add_table(
    ['#', 'Stage', 'Trigger', 'Key Actions & Rules', 'Role Participation'],
    lifecycle_steps,
    col_widths=[0.8, 2.5, 3, 5, 3.5]
)

# -- 1.6 Approval Workflows Summary --
doc.add_heading('1.6 Approval Workflows Summary', level=2)

approval_rows = [
    ['GOM Approval', 'GOM% calculated in Qualification', 'GOM% ≥ auto-approve threshold', 'Reporting Manager', 'Yes — blocks Proposal'],
    ['Discount Approval', 'Discount > 15% AND margin < 20%', 'Otherwise auto-approved', 'Finance Manager', 'No'],
    ['SOW Approval', 'SOW submitted for review', 'N/A — always manual', 'Configurable chain\n(Sales → Presales → Manager\n→ Finance → Legal)', 'No'],
]

add_table(
    ['Approval Type', 'Trigger', 'Auto-Approve Condition', 'Reviewer', 'Stage Gate?'],
    approval_rows,
    col_widths=[2.5, 3.5, 3, 3.5, 2.5]
)

doc.add_paragraph()

# ============== 2. ROLE OVERVIEW ==============
doc.add_heading('2. Role Overview', level=1)

p = doc.add_paragraph()
run = p.add_run('Q-CRM uses a role-based access control (RBAC) system with six system-defined roles. Each role carries a permission set that controls access to features across both the frontend UI and backend API. Users can hold multiple roles and switch between them; the ')
run.font.size = Pt(10)
run = p.add_run('active role')
run.bold = True
run.font.size = Pt(10)
run = p.add_run(' determines current permissions.')
run.font.size = Pt(10)

add_table(
    ['Role', 'Type', 'Target User', 'Summary'],
    [
        ['Admin', 'System', 'IT / System Admin', 'Full system access with wildcard (*) permission'],
        ['Manager', 'System', 'Practice Head / Delivery Manager', 'Full opportunity lifecycle + approvals + analytics export + audit'],
        ['Sales', 'System', 'Sales Executive / Account Manager', 'Pipeline management + sales conversion + lead generation'],
        ['Presales', 'System', 'Presales Consultant / Solution Architect', 'Estimation + presales activities + solution design'],
        ['Management', 'System', 'CXO / Senior Leadership', 'Read-only oversight + approvals + analytics export + audit'],
        ['Read-Only', 'System', 'Stakeholders / External Viewers', 'View-only access across all modules'],
    ],
    col_widths=[2.5, 1.5, 4, 7]
)

# ============== 3. ADMIN ==============
doc.add_heading('3. Admin', level=1)

p = doc.add_paragraph()
run = p.add_run('Persona: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('System Administrator, IT Admin')
run.font.size = Pt(10)
p = doc.add_paragraph()
run = p.add_run('Permission: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('* (wildcard — grants every permission)')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)

admin_sections = {
    '3.1 User Management': [
        ['UC-A01', 'View all users', 'See paginated user list with department, designation, role, status, manager, and mute filters'],
        ['UC-A02', 'Create a new user', 'Add user with name, email, role assignment, and optional department'],
        ['UC-A03', 'Assign/change user roles', 'Multi-select roles for any user from the roles dropdown'],
        ['UC-A04', 'Activate/deactivate users', 'Toggle a user between Active and Inactive status'],
        ['UC-A05', 'Reset user password', 'Set a new password for any non-SSO user'],
        ['UC-A06', 'Assign local password to SSO user', 'Give SSO users a local password for fallback authentication'],
        ['UC-A07', 'Sync users from QPeople HRMS', 'Import/update employee records from the QPeople HR system'],
        ['UC-A08', 'Mute/unmute notifications', 'Toggle email notifications per user'],
        ['UC-A09', 'Reset all user roles', 'Bulk-remove all role assignments (Admin users retain Admin role)'],
    ],
    '3.2 Role Management': [
        ['UC-A10', 'View roles permission matrix', 'See tabular grid of all roles vs. all permission categories'],
        ['UC-A11', 'Edit role permissions', 'Toggle individual permissions on/off for any role'],
        ['UC-A12', 'Create custom role', 'Define new role with name, description, and selected permissions'],
        ['UC-A13', 'Delete custom role', 'Remove non-system roles that have no assigned users'],
        ['UC-A14', 'Reset role defaults', 'Restore all system role permissions to factory defaults'],
        ['UC-A15', 'Add/remove users from roles', 'Manage role membership from the Roles tab'],
    ],
    '3.3 QPeople Integration': [
        ['UC-A16', 'View QPeople role mappings', 'See designation-to-role mapping table'],
        ['UC-A17', 'Create/edit mapping', 'Map a QPeople designation to one or more Q-CRM roles'],
        ['UC-A18', 'Delete mapping', 'Remove a specific designation-to-role mapping'],
        ['UC-A19', 'Reset all mappings', 'Bulk-delete all QPeople role mappings'],
        ['UC-A20', 'Apply mappings', 'Push current mappings to reassign roles for all synced users'],
    ],
    '3.4 Authentication Configuration': [
        ['UC-A21', 'View auth config', 'See current authentication mode and SSO settings'],
        ['UC-A22', 'Switch auth mode', 'Change between Local, SSO, or Hybrid authentication'],
        ['UC-A23', 'Configure SSO settings', 'Set SSO provider, client ID, tenant, redirect URIs'],
    ],
    '3.5 Master Data Administration': [
        ['UC-A24', 'Manage clients', 'Create, edit, delete client records'],
        ['UC-A25', 'Manage regions', 'Create, edit, delete geographic regions'],
        ['UC-A26', 'Manage technologies', 'Create, edit, delete technology/skill tags'],
        ['UC-A27', 'Manage pricing models', 'Create, edit, delete pricing model definitions (T&M, Fixed, etc.)'],
        ['UC-A28', 'Manage project types', 'Create, edit, delete project type categories'],
    ],
    '3.6 Cost & Rate Management': [
        ['UC-A29', 'Manage rate cards', 'Create, edit, delete rate card entries (role/band/rate)'],
        ['UC-A30', 'Manage budget assumptions', 'Configure default budget parameters'],
        ['UC-A31', 'Manage currency rates', 'Add, edit, toggle, delete currency exchange rates'],
        ['UC-A32', 'Sync currency rates', 'Pull latest exchange rates from external source'],
        ['UC-A33', 'Use GOM calculator', 'Run Gross Operating Margin calculations'],
    ],
    '3.7 Email & Notifications': [
        ['UC-A34', 'Manage email templates', 'Create, edit, delete email templates with WYSIWYG editor'],
        ['UC-A35', 'Use template field catalog', 'Insert variables from 15+ data tables (95+ fields)'],
        ['UC-A36', 'Create custom formula fields', 'Define calculated fields using built-in functions'],
        ['UC-A37', 'Preview email templates', 'Live preview with sample data rendering'],
        ['UC-A38', 'Send test emails', 'Dispatch test email to verify template rendering'],
        ['UC-A39', 'Manage notification rules', 'Create, edit, delete trigger rules (opportunity_created, stage_change, data_condition)'],
    ],
    '3.8 SOW Studio Administration': [
        ['UC-A40', 'Manage SOW templates', 'Upload, edit, delete SOW document templates'],
        ['UC-A41', 'Configure template anchors', 'Map content sections to template placeholders'],
        ['UC-A42', 'Manage metadata categories/values', 'Define SOW metadata taxonomies'],
        ['UC-A43', 'Manage static content blocks', 'Create reusable SOW text sections'],
        ['UC-A44', 'Manage clauses', 'Create, edit, delete legal/contractual clauses'],
        ['UC-A45', 'Manage section rules', 'Define conditional section inclusion logic'],
        ['UC-A46', 'Configure approval workflows', 'Set up SOW approval chains and config'],
        ['UC-A47', 'Configure SOW numbering', 'Set document numbering format and sequence'],
    ],
    '3.9 Audit & Compliance': [
        ['UC-A48', 'View audit logs', 'Browse all system audit events with entity/action filters'],
        ['UC-A49', 'Export audit data', 'Download audit log records for compliance review'],
    ],
}

for section_title, rows in admin_sections.items():
    doc.add_heading(section_title, level=3)
    add_table(['#', 'Use Case', 'Details'], rows, col_widths=[1.5, 4, 9.5])

doc.add_heading('3.10 All Operational Features', level=3)
p = doc.add_paragraph()
run = p.add_run('Admin inherits ')
run.font.size = Pt(10)
run = p.add_run('all')
run.bold = True
run.italic = True
run.font.size = Pt(10)
run = p.add_run(' use cases from every other role (Manager, Sales, Presales, Management, Read-Only) due to wildcard permission.')
run.font.size = Pt(10)

# ============== 4. MANAGER ==============
doc.add_heading('4. Manager', level=1)

p = doc.add_paragraph()
run = p.add_run('Persona: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Practice Head, Delivery Manager, Department Lead')
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Permissions: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('dashboard:view, pipeline:view, pipeline:write, presales:view, presales:write, sales:view, sales:write, estimation:manage, approvals:manage, contacts:view, contacts:write, analytics:view, analytics:export, agents:execute, gom:view, leads:manage, resources:manage, settings:view, auditlogs:view')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

manager_sections = {
    '4.1 Dashboard & Pipeline': [
        ['UC-M01', 'View dashboard', 'See high-level metrics, pipeline summary cards, charts'],
        ['UC-M02', 'View opportunity pipeline', 'Browse all opportunities with filters (stage, practice, client, date)'],
        ['UC-M03', 'Create new opportunity', 'Submit a new opportunity with all required fields'],
        ['UC-M04', 'Edit opportunity details', 'Update any field — stage, value, dates, team, practice, etc.'],
        ['UC-M05', 'View opportunity detail', 'Full detail view with timeline, comments, attachments, audit log'],
        ['UC-M06', 'Add comments', 'Post comments on opportunities for team collaboration'],
        ['UC-M07', 'Upload/download attachments', 'Attach supporting documents (RFP, proposal, contracts)'],
        ['UC-M08', 'Delete attachments', 'Remove attached files from opportunities'],
        ['UC-M09', 'View opportunity audit log', 'See change history for any opportunity'],
    ],
    '4.2 Presales & Estimation': [
        ['UC-M10', 'View presales data', 'See solution design, effort estimates, team composition'],
        ['UC-M11', 'Edit presales details', 'Update technical solution, resource plan, effort breakdown'],
        ['UC-M12', 'Manage estimations', 'Create, edit, review effort and cost estimates'],
        ['UC-M13', 'Submit GOM for approval', 'Send Gross Operating Margin calculation for review'],
        ['UC-M14', 'Review GOM approvals', 'Approve or reject GOM submissions from team members'],
    ],
    '4.3 Sales & Conversion': [
        ['UC-M15', 'View sales data', 'See commercial terms, pricing, negotiation status'],
        ['UC-M16', 'Edit sales details', 'Update pricing, discounts, terms, contract values'],
        ['UC-M17', 'Convert opportunity', 'Move opportunity from presales to active sales/won status'],
    ],
    '4.4 Approvals': [
        ['UC-M18', 'Review and approve requests', 'Act on pending approval items (GOM, SOW, stage gates)'],
        ['UC-M19', 'Reject with feedback', 'Decline requests with comments for correction'],
    ],
    '4.5 Analytics & Reporting': [
        ['UC-M20', 'View analytics dashboards', 'Access pipeline analytics, win/loss ratios, trend charts'],
        ['UC-M21', 'Export analytics data', 'Download reports and data extracts in supported formats'],
    ],
    '4.6 Leads & Contacts': [
        ['UC-M22', 'Manage leads', 'Create, update, qualify, and convert leads to opportunities'],
        ['UC-M23', 'View contacts', 'Browse contact directory'],
        ['UC-M24', 'Create/edit contacts', 'Add or update client contact information'],
    ],
    '4.7 Resources & AI': [
        ['UC-M25', 'Manage resources', 'View and manage resource allocation and availability'],
        ['UC-M26', 'Execute AI agents', 'Run AI-powered tasks (chatbot queries, automated analysis)'],
    ],
    '4.8 GOM, Settings & Audit': [
        ['UC-M27', 'Use GOM calculator', 'Calculate Gross Operating Margin for opportunities'],
        ['UC-M28', 'View settings', 'Access read-only view of system configuration'],
        ['UC-M29', 'View audit logs', 'Browse system audit events for oversight and tracking'],
    ],
}

for section_title, rows in manager_sections.items():
    doc.add_heading(section_title, level=3)
    add_table(['#', 'Use Case', 'Details'], rows, col_widths=[1.5, 4, 9.5])

# ============== 5. SALES ==============
doc.add_heading('5. Sales', level=1)

p = doc.add_paragraph()
run = p.add_run('Persona: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Sales Executive, Account Manager, Business Development Manager')
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Permissions: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('dashboard:view, pipeline:view, pipeline:write, presales:view, sales:view, sales:write, contacts:view, contacts:write, analytics:view, agents:execute, gom:view, leads:manage, settings:view')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

sales_sections = {
    '5.1 Dashboard & Pipeline': [
        ['UC-S01', 'View dashboard', 'See personal and team pipeline metrics'],
        ['UC-S02', 'View pipeline', 'Browse opportunities by stage, value, client'],
        ['UC-S03', 'Create opportunity', 'Register new opportunities sourced from leads or direct'],
        ['UC-S04', 'Edit opportunity', 'Update opportunity details, move through stages'],
        ['UC-S05', 'Add comments', 'Collaborate on opportunities via comments'],
        ['UC-S06', 'Manage attachments', 'Upload/download/delete supporting documents'],
        ['UC-S07', 'View audit trail', 'See change history for opportunities'],
    ],
    '5.2 Sales Activities': [
        ['UC-S08', 'View sales data', 'See pricing, commercial terms, contract status'],
        ['UC-S09', 'Edit sales details', 'Update pricing, discounts, win probabilities, close dates'],
        ['UC-S10', 'Convert opportunity', 'Transition opportunity to Won/Active stage'],
    ],
    '5.3 Presales (View Only)': [
        ['UC-S11', 'View presales data', 'See solution details and estimates (read-only)'],
        ['UC-S12', 'View estimations', 'Review effort/cost estimates prepared by presales'],
    ],
    '5.4 Leads & Contacts': [
        ['UC-S13', 'Manage leads', 'Create, qualify, nurture leads and convert to opportunities'],
        ['UC-S14', 'View contacts', 'Access client contact directory'],
        ['UC-S15', 'Create/edit contacts', 'Maintain client stakeholder information'],
    ],
    '5.5 Analytics & AI': [
        ['UC-S16', 'View analytics', 'Access pipeline reports and win/loss analysis'],
        ['UC-S17', 'Execute AI agents', 'Use chatbot and AI-assisted tools'],
        ['UC-S18', 'View GOM calculator', 'Review margin calculations (read-only)'],
        ['UC-S19', 'View settings', 'Access system configuration (read-only)'],
    ],
}

for section_title, rows in sales_sections.items():
    doc.add_heading(section_title, level=3)
    add_table(['#', 'Use Case', 'Details'], rows, col_widths=[1.5, 4, 9.5])

doc.add_heading('5.6 Restrictions', level=3)
restrictions = [
    'Edit presales/estimation data',
    'Approve/reject requests',
    'Export analytics',
    'Manage resources',
    'View audit logs',
    'Access any admin or configuration features',
]
for r in restrictions:
    add_bullet(r, bold_prefix='✗  ')

# ============== 6. PRESALES ==============
doc.add_heading('6. Presales', level=1)

p = doc.add_paragraph()
run = p.add_run('Persona: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Presales Consultant, Solution Architect, Technical Lead')
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Permissions: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('dashboard:view, pipeline:view, presales:view, presales:write, estimation:manage, sales:view, contacts:view, analytics:view, agents:execute, gom:view, settings:view')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

presales_sections = {
    '6.1 Dashboard & Pipeline': [
        ['UC-P01', 'View dashboard', 'See presales workload and pipeline summary'],
        ['UC-P02', 'View pipeline', 'Browse opportunities assigned for presales work'],
        ['UC-P03', 'View opportunity detail', 'Access full opportunity information'],
        ['UC-P04', 'Add comments', 'Collaborate on technical aspects via comments'],
    ],
    '6.2 Presales & Estimation': [
        ['UC-P05', 'Edit presales data', 'Update solution approach, architecture, technical details'],
        ['UC-P06', 'Manage estimations', 'Create and refine effort estimates, resource plans'],
        ['UC-P07', 'Build cost estimates', 'Use rate cards and budget assumptions to build cost models'],
        ['UC-P08', 'Submit GOM for approval', 'Send margin calculation for manager approval'],
        ['UC-P09', 'Review GOM approvals', 'Approve or reject GOM submissions'],
        ['UC-P10', 'Define resource requirements', 'Specify roles, skills, and FTE requirements'],
    ],
    '6.3 Sales & Other (View Only)': [
        ['UC-P11', 'View sales data', 'See commercial terms and pricing decisions (read-only)'],
        ['UC-P12', 'View contacts', 'Browse client contact directory (read-only)'],
        ['UC-P13', 'View analytics', 'Access pipeline and estimation analytics'],
        ['UC-P14', 'Execute AI agents', 'Use chatbot for technical research'],
        ['UC-P15', 'View GOM calculator', 'Review and use margin calculations'],
        ['UC-P16', 'View settings', 'Access system configuration (read-only)'],
    ],
}

for section_title, rows in presales_sections.items():
    doc.add_heading(section_title, level=3)
    add_table(['#', 'Use Case', 'Details'], rows, col_widths=[1.5, 4, 9.5])

doc.add_heading('6.4 Restrictions', level=3)
restrictions = [
    'Create new opportunities or edit pipeline fields (stage, value, dates)',
    'Edit sales data (pricing, discounts) or convert opportunities',
    'Manage leads or create/edit contacts',
    'Export analytics or manage resources',
    'Approve non-GOM requests',
    'View audit logs or access any admin features',
]
for r in restrictions:
    add_bullet(r, bold_prefix='✗  ')

# ============== 7. MANAGEMENT ==============
doc.add_heading('7. Management', level=1)

p = doc.add_paragraph()
run = p.add_run('Persona: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('CXO, VP, Senior Director, Board Member')
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Permissions: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('dashboard:view, pipeline:view, presales:view, sales:view, contacts:view, analytics:view, analytics:export, approvals:manage, auditlogs:view, gom:view, settings:view')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

mgmt_sections = {
    '7.1 Strategic Oversight': [
        ['UC-G01', 'View dashboard', 'See executive-level metrics and KPIs'],
        ['UC-G02', 'View pipeline', 'Browse full opportunity pipeline across all practices'],
        ['UC-G03', 'View opportunity details', 'Drill down into any opportunity for review'],
        ['UC-G04', 'View presales data', 'Review solution approaches and estimates'],
        ['UC-G05', 'View sales data', 'Review commercial terms and revenue projections'],
        ['UC-G06', 'View contacts', 'Access client stakeholder directory'],
    ],
    '7.2 Analytics & Reporting': [
        ['UC-G07', 'View analytics', 'Access comprehensive dashboards, trends, win/loss analysis'],
        ['UC-G08', 'Export analytics', 'Download reports for board presentations and reviews'],
    ],
    '7.3 Approvals & Compliance': [
        ['UC-G09', 'Review approvals', 'Act on escalated approval requests (GOM, SOW, high-value deals)'],
        ['UC-G10', 'Approve/reject with feedback', 'Provide executive approval or request revisions'],
        ['UC-G11', 'View audit logs', 'Review system activity for governance and compliance'],
        ['UC-G12', 'View GOM calculator', 'Review margin calculations'],
        ['UC-G13', 'View settings', 'See system configuration (read-only)'],
    ],
}

for section_title, rows in mgmt_sections.items():
    doc.add_heading(section_title, level=3)
    add_table(['#', 'Use Case', 'Details'], rows, col_widths=[1.5, 4, 9.5])

doc.add_heading('7.4 Restrictions', level=3)
restrictions = [
    'Create or edit opportunities, presales, or sales data',
    'Create or manage estimations, leads, or contacts',
    'Manage resources or execute AI agents',
    'Upload/delete attachments',
    'Access any admin or configuration features',
]
for r in restrictions:
    add_bullet(r, bold_prefix='✗  ')

# ============== 8. READ-ONLY ==============
doc.add_heading('8. Read-Only', level=1)

p = doc.add_paragraph()
run = p.add_run('Persona: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('External Stakeholder, Auditor, Observer, New Employee in Onboarding')
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Permissions: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('dashboard:view, pipeline:view, presales:view, sales:view, contacts:view, analytics:view, gom:view, settings:view')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_heading('8.1 View-Only Access', level=3)
add_table(['#', 'Use Case', 'Details'], [
    ['UC-R01', 'View dashboard', 'See summary metrics and pipeline overview'],
    ['UC-R02', 'View pipeline', 'Browse opportunity list (read-only)'],
    ['UC-R03', 'View opportunity details', 'See full opportunity information, timeline, comments'],
    ['UC-R04', 'View presales data', 'See solution and estimation details'],
    ['UC-R05', 'View sales data', 'See commercial terms and pricing'],
    ['UC-R06', 'View contacts', 'Browse contact directory'],
    ['UC-R07', 'View analytics', 'Access dashboards and reports (no export)'],
    ['UC-R08', 'View GOM calculator', 'See margin calculations'],
    ['UC-R09', 'View settings', 'See system configuration'],
], col_widths=[1.5, 4, 9.5])

doc.add_heading('8.2 Restrictions', level=3)
restrictions = [
    'Create, edit, or delete any records',
    'Add comments or manage attachments',
    'Manage leads or export data',
    'Execute AI agents or approve anything',
    'Access any admin features',
]
for r in restrictions:
    add_bullet(r, bold_prefix='✗  ')

# ============== 9. PERMISSION MATRIX ==============
doc.add_page_break()
doc.add_heading('9. Permission Reference Matrix', level=1)

p = doc.add_paragraph()
run = p.add_run('✅ = Granted    ❌ = Not Granted    * = Can be manually assigned by Admin')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

matrix_rows = [
    ['dashboard:view',   '✅', '✅', '✅', '✅', '✅', '✅'],
    ['pipeline:view',    '✅', '✅', '✅', '✅', '✅', '✅'],
    ['pipeline:write',   '✅', '✅', '✅', '❌', '❌', '❌'],
    ['presales:view',    '✅', '✅', '✅', '✅', '✅', '✅'],
    ['presales:write',   '✅', '✅', '❌', '✅', '❌', '❌'],
    ['estimation:manage','✅', '✅', '❌', '✅', '❌', '❌'],
    ['sales:view',       '✅', '✅', '✅', '✅', '✅', '✅'],
    ['sales:write',      '✅', '✅', '✅', '❌', '❌', '❌'],
    ['approvals:manage', '✅', '✅', '❌', '❌', '✅', '❌'],
    ['contacts:view',    '✅', '✅', '✅', '✅', '✅', '✅'],
    ['contacts:write',   '✅', '✅', '✅', '❌', '❌', '❌'],
    ['analytics:view',   '✅', '✅', '✅', '✅', '✅', '✅'],
    ['analytics:export', '✅', '✅', '❌', '❌', '✅', '❌'],
    ['agents:execute',   '✅', '✅', '✅', '✅', '❌', '❌'],
    ['gom:view',         '✅', '✅', '✅', '✅', '✅', '✅'],
    ['leads:manage',     '✅', '✅', '✅', '❌', '❌', '❌'],
    ['resources:manage', '✅', '✅', '❌', '❌', '❌', '❌'],
    ['settings:view',    '✅', '✅', '✅', '✅', '✅', '✅'],
    ['settings:manage',  '✅', '❌', '❌', '❌', '❌', '❌'],
    ['users:manage',     '✅', '❌', '❌', '❌', '❌', '❌'],
    ['roles:manage',     '✅', '❌', '❌', '❌', '❌', '❌'],
    ['metadata:manage',  '✅', '❌', '❌', '❌', '❌', '❌'],
    ['costcard:manage',  '✅', '❌', '❌', '❌', '❌', '❌'],
    ['auditlogs:view',   '✅', '✅', '❌', '❌', '✅', '❌'],
    ['sow:view',         '✅', '❌*', '❌*', '❌*', '❌', '❌'],
    ['sow:write',        '✅', '❌*', '❌', '❌*', '❌', '❌'],
    ['sow:admin',        '✅', '❌', '❌', '❌', '❌', '❌'],
]

add_table(
    ['Permission', 'Admin', 'Manager', 'Sales', 'Presales', 'Mgmt', 'Read-Only'],
    matrix_rows,
    col_widths=[3.5, 1.5, 1.8, 1.5, 1.8, 1.5, 1.8],
    header_color='2c3e50'
)

# ============== NOTES ==============
doc.add_heading('Notes', level=1)

notes = [
    ('Custom Roles — ', 'Admins can create additional roles with any combination of permissions beyond these six system roles.'),
    ('Multiple Roles — ', 'Users can be assigned multiple roles and switch between them; the active role determines current access.'),
    ('Wildcard — ', 'Only Admin has * permission; this automatically grants access to any current and future features.'),
    ('SSO Users — ', 'Users authenticating via SSO get roles assigned through QPeople designation mappings or manual admin assignment.'),
    ('Notification Muting — ', 'Any user can be muted by an Admin regardless of role; muted users do not receive system email notifications.'),
]
for bold_part, rest in notes:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_part)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run = p.add_run(rest)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

# ============== FOOTER ==============
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 30)
run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Q-CRM Role Use Cases  •  https://qcrm.qbadvisory.com  •  QB Advisory')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Calibri'

# Save
output_path = r'd:\Opportunity\Jaydeep_work\QCRM_Role_Use_Cases.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
